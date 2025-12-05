# exporters.py

import pandas as pd
from io import BytesIO
from openpyxl.styles import PatternFill, Font
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT


# ----------------------------
# Excel Export
# ----------------------------
def export_excel(df):
    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Upah")
        sheet = writer.sheets["Upah"]

        # Auto column width
        for i, col in enumerate(df.columns):
            width = max(df[col].astype(str).map(len).max(), len(col)) + 2
            sheet.column_dimensions[chr(65 + i)].width = width

        # Header bold
        for cell in sheet[1]:
            cell.font = Font(bold=True)

        # Highlight "Total" row
        for row_idx, row in df.iterrows():
            if row["Bulan"] == "Total":
                for col_idx in range(len(df.columns)):
                    cell = sheet.cell(row_idx + 2, col_idx + 1)
                    cell.fill = PatternFill(start_color="FFF3B0", fill_type="solid")
                    cell.font = Font(bold=True)

    return output.getvalue()


# ----------------------------
# Word Export
# ----------------------------
def export_word(df):
    doc = Document()
    doc.add_heading("Lampiran Kekurangan Upah", level=1)

    table = doc.add_table(rows=1, cols=len(df.columns), style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    header_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        header_cells[i].text = col
        header_cells[i].paragraphs[0].runs[0].font.bold = True

    # Data rows
    for idx, row in df.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)

            # Highlight total row
            if row["Bulan"] == "Total":
                run = cells[i].paragraphs[0].runs[0]
                run.font.bold = True

                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "FFF176")
                cells[i]._tc.get_or_add_tcPr().append(shading)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
